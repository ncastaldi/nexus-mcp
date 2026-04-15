"""Multi-format report renderer for Nexus MCP.

Converts a markdown string to the requested output format.
All non-markdown renderers have optional dependencies guarded with
clear ImportError messages — the server starts fine without them.

Formats
-------
md   — plain markdown (no extra deps)
html — standalone HTML with Wheels-branded CSS
         requires: pip install markdown
pdf  — PDF generated from the branded HTML
         requires: pip install weasyprint markdown
         Windows note: GTK3 runtime DLLs are also required
         https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#windows
docx — Word document
         requires: pip install python-docx
"""

from __future__ import annotations

import io
from enum import Enum


# ── Brand / style constants ───────────────────────────────────────────────────

_BRAND_PRIMARY = "#003865"   # Wheels dark blue
_BRAND_ACCENT  = "#0066cc"   # Wheels mid blue
_BRAND_ROW_ALT = "#f5f8fc"   # table row stripe

_CSS = f"""
body {{
    font-family: 'Segoe UI', Arial, sans-serif;
    margin: 40px auto;
    max-width: 960px;
    color: #222;
    line-height: 1.55;
}}
h1 {{
    color: {_BRAND_PRIMARY};
    border-bottom: 3px solid {_BRAND_PRIMARY};
    padding-bottom: 8px;
    margin-bottom: 4px;
}}
h2 {{
    color: {_BRAND_ACCENT};
    margin-top: 32px;
}}
h3 {{ color: #444; }}
table {{
    border-collapse: collapse;
    width: 100%;
    margin: 16px 0;
    font-size: 0.92em;
}}
th {{
    background: {_BRAND_PRIMARY};
    color: #fff;
    padding: 8px 12px;
    text-align: left;
    font-weight: 600;
}}
td {{
    padding: 7px 12px;
    border-bottom: 1px solid #ddd;
}}
tr:nth-child(even) td {{
    background: {_BRAND_ROW_ALT};
}}
code, pre {{
    background: #f0f0f0;
    padding: 2px 6px;
    border-radius: 3px;
    font-size: 0.9em;
    font-family: 'Consolas', 'Courier New', monospace;
}}
ul, ol {{ padding-left: 1.4em; }}
li {{ margin-bottom: 4px; }}
.meta {{
    color: #666;
    font-size: 0.88em;
    margin-bottom: 20px;
}}
"""

_HTML_WRAPPER = """\
<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{title}</title>
  <style>{css}</style>
</head>
<body>
{body}
</body>
</html>
"""


# ── Public API ────────────────────────────────────────────────────────────────

class ReportFormat(str, Enum):
    MD   = "md"
    HTML = "html"
    PDF  = "pdf"
    DOCX = "docx"

    @property
    def extension(self) -> str:
        return self.value

    @property
    def is_binary(self) -> bool:
        return self in (ReportFormat.PDF, ReportFormat.DOCX)


def render(
    markdown_content: str,
    fmt: ReportFormat,
    title: str = "Report",
) -> bytes | str:
    """Render *markdown_content* to the requested format.

    Returns:
        str  for MD and HTML
        bytes for PDF and DOCX
    """
    if fmt == ReportFormat.MD:
        return markdown_content

    html = _md_to_html(markdown_content, title)

    if fmt == ReportFormat.HTML:
        return html

    if fmt == ReportFormat.PDF:
        return _html_to_pdf(html)

    if fmt == ReportFormat.DOCX:
        return _md_to_docx(markdown_content, title)

    raise ValueError(f"Unsupported format: {fmt!r}")


# ── Internal renderers ────────────────────────────────────────────────────────

def _md_to_html(md: str, title: str) -> str:
    """Convert markdown string to a full, styled HTML document."""
    try:
        import markdown as _md_lib
    except ImportError as exc:
        raise ImportError(
            "HTML/PDF output requires the 'markdown' package.\n"
            "  pip install markdown\n"
            "  or: pip install 'nexus-mcp[report]'"
        ) from exc

    body = _md_lib.markdown(
        md,
        extensions=["tables", "fenced_code", "nl2br"],
    )
    return _HTML_WRAPPER.format(title=title, css=_CSS, body=body)


def _html_to_pdf(html: str) -> bytes:
    """Convert a full HTML string to PDF bytes using WeasyPrint."""
    try:
        from weasyprint import HTML as _HTML
    except ImportError as exc:
        raise ImportError(
            "PDF output requires WeasyPrint.\n"
            "  pip install weasyprint\n"
            "  or: pip install 'nexus-mcp[report]'\n"
            "\n"
            "Windows: GTK3 runtime DLLs are also required:\n"
            "  https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#windows"
        ) from exc

    return _HTML(string=html).write_pdf()


def _md_to_docx(md: str, title: str) -> bytes:
    """Convert markdown to a .docx document using python-docx.

    Handles: headings (#/##/###), bullet lists (- …), pipe tables,
    key-value pairs (- Key: Value), and plain paragraphs.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as exc:
        raise ImportError(
            "DOCX output requires python-docx.\n"
            "  pip install python-docx\n"
            "  or: pip install 'nexus-mcp[report]'"
        ) from exc

    doc = Document()

    # ── document title ────────────────────────────────────────────────────────
    heading = doc.add_heading(title, level=0)
    heading.runs[0].font.color.rgb = RGBColor(0x00, 0x38, 0x65)

    lines = md.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip blank lines
        if not line.strip():
            i += 1
            continue

        # Heading 1
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        # Heading 2
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # Heading 3
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Bullet
        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            i += 1
            continue

        # Pipe table: collect header + separator + data rows
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|---"):
            headers = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2  # skip separator row
            data_rows: list[list[str]] = []
            while i < len(lines) and lines[i].startswith("|"):
                data_rows.append(
                    [c.strip() for c in lines[i].strip().strip("|").split("|")]
                )
                i += 1

            ncols = len(headers)
            t = doc.add_table(rows=1 + len(data_rows), cols=ncols)
            t.style = "Table Grid"

            # Header row — bold + brand colour fill
            hdr_row = t.rows[0]
            for j, cell_text in enumerate(headers):
                cell = hdr_row.cells[j]
                cell.text = cell_text
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                # Fill background
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "003865")
                tc_pr.append(shd)

            for r_idx, row_data in enumerate(data_rows):
                for j, cell_text in enumerate(row_data):
                    if j < ncols:
                        t.rows[r_idx + 1].cells[j].text = cell_text

            continue

        # Fallback: plain paragraph
        doc.add_paragraph(line)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
